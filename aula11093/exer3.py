from docx import Document

doc = Document()

num = int(input("Gerar tabuada de? "))

doc.add_heading('Tabuada de Multiplicação de {}'.format(num), 0)

tab = doc.add_table(rows=1, cols=5)
tab.style="Medium Grid 1 Accent 2"
cels = tab.rows[0].cells
cels[0].text = 'Número'
cels[1].text = "x"
cels[2].text = "Fator"
cels[3].text = "="
cels[4].text = "Resultado"

for fator in range(0, 11):
  mult = num * fator
  dados = tab.add_row().cells
  dados[0].text = str(num)
  dados[1].text = "x"
  dados[2].text = str(fator)
  dados[3].text = "="
  dados[4].text = str(mult)

doc.save("c:/temp/tabuada.docx")


